try:
    import argparse
    from fake_headers import Headers
    import requests
    import json
    import docx
except ModuleNotFoundError:
    print("Please download dependencies from requirement.txt")
except Exception as ex:
    print(ex)


class Pinterest:
    '''This class scraps pinterest and returns a dict containing all user data'''
    @staticmethod
    def _generate_url(username):
        return "https://pinterest.com/resource/UserResource/get/?source_url=%25{}%2F&data=%7B%22options%22%3A%7B%22field_set_key%22%3A%22profile%22%2C%22username%22%3A%22{}%22%2C%22is_mobile_fork%22%3Atrue%7D%2C%22context%22%3A%7B%7D%7D&_=1640428319046".format(username, username)

    @staticmethod
    def _make_request(url):
        headers = Headers().generate()
        response = requests.get(url, headers=headers)
        return response

    @staticmethod
    def scrap(username):
        try:

            try:
                url = Pinterest._generate_url(username)
                response = Pinterest._make_request(url)
                if response.status_code == 200:
                    response = response.json()
                else:
                    print("Failed to get Data!")
                    exit()
            except Exception as ex:
                print("Error", ex)
                exit()

            json_data = response
            print(json_data)
            data = json_data['resource_response']['data']

            return json.dumps(data)
        except Exception as ex:
            print(ex)


def scrapPinterest(username):
    
    print(Pinterest.scrap(username))

    doc = docx.Document()
  
    # Add a Title to the document
    doc.add_heading('Details', 0)
    data = json.loads(Pinterest.scrap(username))
    print(type(data))
    print(len(data))
    table = doc.add_table(rows=1, cols=2)
  
    # Adding heading in the 1st row of the table
    row = table.rows[0].cells
    row[0].text = 'Id'
    row[1].text = 'Value'

    for id in data.keys():
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(id)
        row[1].text = str(data[id])
    doc.save('pinterest.docx')


# last updated - 8th July,2022
