from linkedin_api import Linkedin
import docx

def scrapLinkedin(un):
    api = Linkedin('policeHackathon23@gmail.com', 'Police@123')
    profile = api.get_profile(un)
    print(profile)

    doc = docx.Document()
    
    # Add a Title to the document
    doc.add_heading('Details', 0)

    table = doc.add_table(rows=1, cols=2)

    row = table.add_row().cells
    row[0].text = "Summary"
    row[1].text = profile['summary']
    row = table.add_row().cells
    row[0].text = "Name of the Industry"
    row[1].text = profile['industryName']
    row = table.add_row().cells
    row[0].text = "First Name"
    row[1].text = profile['firstName']
    row = table.add_row().cells
    row[0].text = "Last Name"
    row[1].text = profile['lastName']
    row = table.add_row().cells
    row[0].text = "Geo Location"
    row[1].text = profile['geoLocationName']
    row = table.add_row().cells
    row[0].text = "Country"
    row[1].text = profile['geoCountryName']
    row = table.add_row().cells
    row[0].text = "Headline"
    row[1].text = profile['headline']
    row = table.add_row().cells
    row[0].text = "Experiences"
    for exp in profile['experience']:
        row[1].text = row[1].text + exp['companyName'] + " - " + exp['title'] + "; "
    row = table.add_row().cells
    row[0].text = "Education"
    for edu in profile['education']:
        row[1].text = row[1].text + edu['schoolName'] + "; "

    doc.save('linkedin.docx')