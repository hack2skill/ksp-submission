try:
    import argparse
    import requests
    import json
    import pandas as pd
    import docx
    import os
except ModuleNotFoundError:
    print("Please download dependencies from requirement.txt")
except Exception as ex:
    print(ex)


class Medium:

    @staticmethod
    def build_payload(username):
        query = ''
        print("XXX")
        module_dir = os.path.dirname(__file__)  # get current directory
        file_path = os.path.join(module_dir, 'medium_graphql_query.graphql')
        print(file_path)
        with open(file_path, 'r', encoding='utf-8') as file:
            query = file.read()
        json_data = [
            {
                'operationName': 'UserProfileQuery',
                'variables': {
                    'includeDistributedResponses': True,
                    'id': None,
                    'username': '@{}'.format(username),
                    'homepagePostsLimit': 1,
                },
                'query': query
            },
        ]
        return json_data

    @staticmethod
    def make_request(URL, json_data):
        try:
            response = requests.post(URL, json=json_data)
            if response.status_code == 200:
                return response.json()
        except Exception as ex:
            print('Error at Make Request: {}'.format(ex))

    @staticmethod
    def scrap(username):
        """scrap medium's profile"""
        try:
            URL = "https://medium.com/_/graphql"
            payload = Medium.build_payload(username)
            response = Medium.make_request(URL, payload)
            print("Hello World??")
            if response:
                print("Hello World?")
                return json.dumps(response)
            else:
                print("Failed to make response!")
        except Exception as ex:
            return {"error": ex}


def scrapMedium(username):
    
    print(type(Medium.scrap(username)))

    doc = docx.Document()
  
    # Add a Title to the document
    doc.add_heading('Details', 0)
    data = json.loads(Medium.scrap(username))
    print("HelloWorld")
    print(type(data))
    print("hELLOwORLD")
    print(len(data))
    table = doc.add_table(rows=1, cols=2)
  
    # Adding heading in the 1st row of the table
    row = table.rows[0].cells
    row[0].text = 'Id'
    row[1].text = 'Value'
    
    # Adding data from the list to the table
    """df = pd.DataFrame.from_dict({(i,j): data[0][i][j] 
                           for i in data[0].keys() 
                           for j in data[0][i].keys()},
                       orient='index')"""
    """for id in data[0].keys():
        print(len(data[0]))
        print(type(data[0][id]['userResult']))
        temp = data[0][id]['userResult']
        for id2 in temp.keys():
        # Adding a row and then adding data in it.
            row = table.add_row().cells
            # Converting id to string as table can only take string input
            row[0].text = str(id2)
            row[1].text = str(temp[id2])"""
    temp = data[0]['data']['userResult']
    row = table.add_row().cells
    row[0].text = str("Name")
    row[1].text = str(temp['name'])
    row = table.add_row().cells
    row[0].text = str("Username")
    row[1].text = str(temp['username'])
    row = table.add_row().cells
    row[0].text = str("Domain")
    row[1].text = str(temp['customDomainState']['live']['domain'])
    row = table.add_row().cells
    row[0].text = str("Status")
    row[1].text = str(temp['customDomainState']['live']['status'])
    row = table.add_row().cells
    row[0].text = str("Follower count")
    row[1].text = str(temp['socialStats']['followerCount'])
    row = table.add_row().cells
    row[0].text = str("Name")
    row[1].text = str(temp['newsletterV3']['user']['name'])
    row = table.add_row().cells
    row[0].text = str("PromoHeadline")
    row[1].text = str(temp['newsletterV3']['promoHeadline'])
    row = table.add_row().cells
    row[0].text = str("PromoBody")
    row[1].text = str(temp['newsletterV3']['promoBody'])
    row = table.add_row().cells
    row[0].text = str("About")
    row[1].text = str(temp['about'])
    row = table.add_row().cells
    row[0].text = str("Bio")
    row[1].text = str(temp['bio'])
    row = table.add_row().cells
    row[0].text = str("Has Completed profile")
    row[1].text = str(temp['hasCompletedProfile'])
    row = table.add_row().cells
    row[0].text = str("Twitter screen name")
    row[1].text = str(temp['twitterScreenName'])
    row = table.add_row().cells
    row[0].text = str("is Suspended")
    row[1].text = str(temp['isSuspended'])
    # Now save the document to a location
    doc.save('medium.docx')
# last modified on : 24th October, 2022
