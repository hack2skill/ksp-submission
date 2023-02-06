import docx
import json
from .medium_json import Medium_json
from .twitter_json import Twitter_json
from .instagram_json import Insta_json
from .github_json import Github_json
from .pinterest_json import Pinterest_json
from .quora_json import Quora_json
from .linkedin_json import LinkedIn_json
from FacebookPostsScraper.facebook_json import Facebook_json
def finaldata(r):
    print(r)
    if(r[0]!=''):
        fb_json = Facebook_json(r[0])
    if(r[1]!=''):
        Twt_json = Twitter_json(r[1])
    if(r[2]!=''):
        Gh_json = Github_json(r[2])
    if(r[3]!=''):
        Ig_json = Insta_json(r[3])
    if(r[4]!=''):
        li_json = LinkedIn_json(r[4])
    if(r[5]!=''):
        Q_json = Quora_json(r[5])
    if(r[6]!=''):
        Md_json = Medium_json(r[6])
    if(r[7]!=''):
        P_json = Pinterest_json(r[7])

    doc = docx.Document()

    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0].cells
    row[0].text = 'Id'
    row[1].text = 'Value'

    if(r[3]!=''):
        row = table.add_row().cells
        row[0].text = "Full Name"
        fullName = Ig_json['full_name']
        row[1].text = fullName
        row = table.add_row().cells
        row[0].text = "Instagram username"
        instagram_username = Ig_json['username']
        row[1].text = instagram_username
        row[0].text = "Instagram Biography"
        instagram_bio = Ig_json['biography']
        row[1].text = instagram_bio
        row[0].text = "Instagram profile picture URL"
        instagram_pfp = Ig_json['profile_pic_url_hd']
        row[1].text = instagram_pfp

    if(r[4]!=''):
        row = table.add_row().cells
        row[0].text = "Linked Summary"
        li_summary = li_json['summary']
        row[1].text = li_summary
        row = table.add_row().cells
        row[0].text = "Industry Name"
        li_compName = li_json['industryName']
        row[1].text = li_compName
        row = table.add_row().cells
        row[0].text = "LinkedIn Profile Headline"
        li_headline = li_json['headline']
        row[1].text = li_headline
        row = table.add_row().cells
        row[0].text = "Experiences"
        for exp in li_json['experience']:
                row[1].text = row[1].text + exp['companyName'] + " - " + exp['title'] + "; "
                row = table.add_row().cells
                row[0].text = "Education"
        for edu in li_json['education']:
                row[1].text = row[1].text + edu['schoolname'] + "; "
    if(r[1]!=''):
        for id in Twt_json.keys():
                row = table.add_row().cells
                row[0].text = str(id)
                row[1].text = str(Twt_json[id])
    if(r[6]!=''):
        temp = Md_json[0]['data']['userResult']
        row = table.add_row().cells
        row[0].text = str(" Medium Username")
        row[1].text = str(temp['username'])
        row = table.add_row().cells
        row[0].text = str("Medium Follower count")
        row[1].text = str(temp['socialStats']['followerCount'])
        row = table.add_row().cells
        row[0].text = str("Name")
        row[1].text = str(temp['newsletterV3']['user']['name'])
        row = table.add_row().cells
        row[0].text = str("PromoHeadline")
        row[1].text = str(temp['newsletterV3']['promoHeadline'])
        row = table.add_row().cells
        row[0].text = str("PromoBody")
        row[1].text = str(temp['newsletterV3']['promoBody'])
        row = table.add_row().cells
        row[0].text = str("About")
        row[1].text = str(temp['about'])
        row = table.add_row().cells
        row[0].text = str("Medium Bio")
        row[1].text = str(temp['bio'])
        row = table.add_row().cells
        row[0].text = str("Twitter screen name")
        row[1].text = str(temp['twitterScreenName'])
    if(r[5]!=''):
        for id in Q_json.keys():
                # Adding a row and then adding data in it.
                row = table.add_row().cells
                # Converting id to string as table can only take string input
                row[0].text = str(id)
                row[1].text = str(Q_json[id])
    if(r[2]!=''):
        for id in Gh_json.keys():
                # Adding a row and then adding data in it.
                row = table.add_row().cells
                # Converting id to string as table can only take string input
                row[0].text = str(id)
                row[1].text = str(Gh_json[id])
    if(r[7]!=''):
        for id in P_json.keys():
                # Adding a row and then adding data in it.
                row = table.add_row().cells
                # Converting id to string as table can only take string input
                row[0].text = str(id)
                row[1].text = str(P_json[id])
    if(r[0]!=''):

        row = table.add_row().cells
        row[0].text = "Facebook user ID"
        fb_id = fb_json['User ID']
        row[1].text = fb_id
        row = table.add_row().cells
        row[0].text = "Facebook username"
        fb_username = fb_json['Username']
        row[1].text = fb_username
        row = table.add_row().cells
        row[0].text = "Facebook user URL"
        fb_user_url = fb_json['User URL']
        row[1].text = fb_user_url
        row = table.add_row().cells
        row[0].text = "Facbook cover photo"
        fb_cover_pic = fb_json['Cover Photo']
        row[1].text = fb_cover_pic
        row = table.add_row().cells
        row[0].text = "Facebook profile picture"
        fb_profile_pic = fb_json['Profile Picture']
        row[1].text = fb_profile_pic
        row[0].text = "user category"
        fb_category = fb_json['Category']
        row[1].text = fb_category

    
        '''table1 = doc.add_table(rows=1, cols=4)
        row=table1.add_row().cells
        row[0].text = 'Date published'
        row[1].text = 'Description'
        row[2].text = 'Images'
        row[3].text = 'Post URL'''
        for d in fb_json['posts']:
                row=table.add_row().cells
                row[0].text='Date Published'
                row[1].text = d['Date published']
                row=table.add_row().cells
                row[0].text='Description'
                row[1].text = d['Description']
                row=table.add_row().cells
                row[0].text='Images'
                row[1].text = d['Images']
                row=table.add_row().cells
                row[0].text='Post URL'
                row[1].text = d['Post URL']
    doc.save("CombinedReport.docx")