try:
    import argparse
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options as ChromeOptions
    from selenium.webdriver.firefox.options import Options as FirefoxOptions
    from fake_headers import Headers
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from webdriver_manager.chrome import ChromeDriverManager
    from webdriver_manager.firefox import GeckoDriverManager
    import json
    import docx
except ModuleNotFoundError:
    print("Please download dependencies from requirement.txt")

class Quora:
    @staticmethod
    def init_driver(browser_name:str):
        def set_properties(browser_option):
            ua = Headers().generate()      #fake user agent
            browser_option.add_argument('--headless')
            browser_option.add_argument('--disable-extensions')
            browser_option.add_argument('--incognito')
            browser_option.add_argument('--disable-gpu')
            browser_option.add_argument('--log-level=3')
            browser_option.add_argument(f'user-agent={ua}')
            browser_option.add_argument('--disable-notifications')
            browser_option.add_argument('--disable-popup-blocking')
            return browser_option
        try:
            browser_name = browser_name.strip().title()


            #automating and opening URL in headless browser
            if browser_name.lower() == "chrome":
                browser_option = ChromeOptions()
                browser_option = set_properties(browser_option)
                driver = webdriver.Chrome(ChromeDriverManager().install(),options=browser_option) #chromedriver's path in first argument
            elif browser_name.lower() == "firefox":
                browser_option = FirefoxOptions()
                browser_option = set_properties(browser_option)
                driver = webdriver.Firefox(executable_path=GeckoDriverManager().install(),options=browser_option)
            else:
                driver = "Browser Not Supported!"
            return driver
        except Exception as ex:
            print(ex)
    @staticmethod
    def scrap(username,browser_name):
        try:
            URL = 'https://quora.com/profile/{}'.format(username)
            try:

                driver = Quora.init_driver(browser_name)
                driver.get(URL)
            except AttributeError:
                print("Driver is not set")
                exit()



            wait = WebDriverWait(driver, 10)
            element = wait.until(EC.title_contains('Quora'))

            try:
              name = driver.find_element_by_css_selector("div.q-text.qu-bold")
            except Exception as ex:
              print(ex)
            name = name.text if type(name) is not str else ""

            try:
                profession = driver.find_element_by_css_selector(".q-text.qu-wordBreak--break-word")
            except:
                profession = ""
            profession = profession.text.strip() if type(profession) is not str else ""
            try:
                profile_image = driver.find_element_by_css_selector("img.q-image.qu-display--block")
            except:
                profile_image = ""
            try:
                driver.find_element_by_css_selector("div.qt_read_more").click()
            except Exception as ex:
                print(ex)
                pass

            try:
                details = driver.find_elements_by_css_selector('div.q-box.qu-overflowX--hidden.qu-whiteSpace--nowrap')
                details_text = details[0].text.split('\n')
                answers_count = [text.split(' ')[0] for text in details_text if "Answer" in text]
                answers_count = answers_count[0] if len(answers_count) > 0 else ""
                questions = [text.split(' ')[0] for text in details_text if "Question" in text ]
                questions = questions[0] if len(questions) > 0 else ""
                shares = [text.split(' ')[0]
                          for text in details_text if "share" in text ]
                shares = shares[0] if len(shares) > 0 else ""
                posts = [text.split(' ')[0] for text in details_text if "Posts" in text ]
                posts = posts[0] if len(posts) > 0 else ""
                follwing_follower_element = driver.find_element_by_css_selector(
                    '.q-flex.qu-flexDirection--column.qu-mt--tiny')
                elements = follwing_follower_element.find_elements_by_css_selector(
                    '.CssComponent-sc-1oskqb9-0.AbstractSeparatedItems___StyledCssComponent-sc-46kfvf-0')
                followers = [element.text.split(' ')[0] for element in elements if "follower" in element.text ]
                followers = followers[0] if len(followers) > 0 else ""

                followings = [element.text.split(' ')[0] for element in elements if "following" in element.text ]
                followings = followings[0] if len(followings) > 0 else ""
            except Exception as ex:
                print(ex)
                questions = shares = posts = followers = answers_count = ''

            bio_text = ""
            try:
                bio = driver.find_elements_by_css_selector("p.q-text")
                for p_tag in bio:
                  bio_text += p_tag.text
            except:
                bio = ""
            bio = bio_text if type(bio) is not str else ""
            profile_data = {
                'name'  :name,
                'profession' : profession,
                'profile_image' : profile_image.get_attribute("src"),
                'bio' : bio,
                'answers_count' : answers_count,
                'questions_count' : questions,
                'shares' : shares,
                'posts' : posts,
                'followers' : followers,
                'following' : followings,
            }
            driver.close()
            driver.quit()
            return json.dumps(profile_data)
        except Exception as ex:
            driver.close()
            driver.quit()
            print(ex)

def scrapQuora(username):
    print(Quora.scrap(username,"Chrome"))

    doc = docx.Document()
  
    # Add a Title to the document
    doc.add_heading('Details', 0)
    data = json.loads(Quora.scrap(username,"chrome"))
    print(type(data))
    print(len(data))
    table = doc.add_table(rows=1, cols=2)
  
    # Adding heading in the 1st row of the table
    row = table.rows[0].cells
    row[0].text = 'Id'
    row[1].text = 'Value'

    for id in data.keys():
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(id)
        row[1].text = str(data[id])
    doc.save('details.docx')

#last updated - 02nd October, 2022