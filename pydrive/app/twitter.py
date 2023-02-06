try:
    import requests
    import argparse
    from fake_headers import Headers
    import json
    import docx
except ModuleNotFoundError:
    print("Please download dependencies from requirement.txt")
except Exception as ex:
    print(ex)

AUTHORIZATION_KEY = 'Bearer AAAAAAAAAAAAAAAAAAAAANRILgAAAAAAnNwIzUejRCOuH5E6I8xnZz4puTs%3D1Zv7ttfk8LF81IUq16cHjhLTvJu4FA33AGWWjCpTnA'


class Twitter:

    @staticmethod
    def find_x_guest_token():
        try:
            headers = {
                'authorization': AUTHORIZATION_KEY,
            }
            response = requests.post(
                'https://api.twitter.com/1.1/guest/activate.json', headers=headers)
            return response.json()['guest_token']
        except Exception as ex:
            print("Error at find_x_guest_token: {}".format(ex))

    @staticmethod
    def make_http_request(URL, headers):
        try:
            response = requests.get(URL, headers=headers)
            if response and response.status_code == 200:
                return response.json()
        except Exception as ex:
            print("Error at make_http_request: {}".format(ex))

    @staticmethod
    def build_headers(x_guest_token, authorization_key):
        headers = {
            'authority': 'twitter.com',
            'accept': '*/*',
            'accept-language': 'en-US,en;q=0.9',
            'authorization': authorization_key,
            'sec-fetch-dest': 'empty',
            'sec-fetch-mode': 'cors',
            'sec-fetch-site': 'same-origin',
            'user-agent': Headers().generate()['User-Agent'],
            'x-guest-token': x_guest_token,
            'x-twitter-active-user': 'yes',
            'x-twitter-client-language': 'en',
        }
        return headers

    @staticmethod
    def scrap(username):
        try:
            # generating URL according to the username
            guest_token = Twitter.find_x_guest_token()
            headers = Twitter.build_headers(guest_token, AUTHORIZATION_KEY)
            response = Twitter.make_http_request(
                "https://api.twitter.com/1.1/users/show.json?screen_name={}".format(username),
                headers=headers)
            if response:
              return json.dumps(response)
            else:
              print("Failed to make Request!")
        except Exception as ex:
            print(ex)


def scrapTwitter(username):
    print(Twitter.scrap(username))

    doc = docx.Document()
  
    # Add a Title to the document
    doc.add_heading('Details', 0)
    data = json.loads(Twitter.scrap(username))
    print(type(data))
    print(len(data))
    table = doc.add_table(rows=1, cols=2)
  
    # Adding heading in the 1st row of the table
    row = table.rows[0].cells
    row[0].text = 'Id'
    row[1].text = 'Value'

    for id in data.keys():
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(id)
        row[1].text = str(data[id])
    doc.save('twitter.docx')
# last updated - 24th October, 2022
    return Twitter.scrap(username)