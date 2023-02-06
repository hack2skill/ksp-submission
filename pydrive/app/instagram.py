try:
    import argparse
    import json
    import requests
    from fake_headers import Headers
    import docx
except ModuleNotFoundError:
    print("Please download dependencies from requirement.txt")
except Exception as ex:
    print(ex)


'''can scrap only public instagram accounts'''


class Instagram:
    @staticmethod
    def build_param(username):
        params = {
            'username': username,
        }
        return params

    @staticmethod
    def build_headers(username):
        return {
            'authority': 'www.instagram.com',
            'accept': '*/*',
            'accept-language': 'en-US,en;q=0.9',
            'referer': f'https://www.instagram.com/{username}/',
            'sec-ch-prefers-color-scheme': 'dark',
            'sec-ch-ua': '"Not?A_Brand";v="8", "Chromium";v="108", "Microsoft Edge";v="108"',
            'sec-fetch-dest': 'empty',
            'sec-fetch-mode': 'cors',
            'sec-fetch-site': 'same-origin',
            'user-agent': Headers().generate()['User-Agent'],
            'x-asbd-id': '198387',
            'x-csrftoken': 'VUm8uVUz0h2Y2CO1SwGgVAG3jQixNBmg',
            'x-ig-app-id': '936619743392459',
            'x-ig-www-claim': '0',
            'x-requested-with': 'XMLHttpRequest',
        }

    @staticmethod
    def make_request(url, params, headers, proxy=None):
        response = None
        if proxy:
            proxy_dict = {
                'http': f'http://{proxy}',
                'https': f'http://{proxy}'
            }
            response = requests.get(
                url, headers=headers, params=params, proxies=proxy_dict)
        else:
            response = requests.get(
                url, headers=headers, params=params)
        return response

    @staticmethod
    def scrap(username, proxy = None):
        try:
            headers = Instagram.build_headers(username)
            params = Instagram.build_param(username)
            response = Instagram.make_request('https://www.instagram.com/api/v1/users/web_profile_info/',
            headers=headers, params=params, proxy=proxy)
            if response.status_code == 200:
                profile_data = response.json()['data']['user']
                return json.dumps(profile_data)
            else:
                print('Error : ', response.status_code, response.text)
        except Exception as ex:
            print(ex)


def scrapInstagram(un):
    #parser = argparse.ArgumentParser()
    #parser.add_argument("username", help="username to search")
    #parser.add_argument("--proxy", help="Proxy to use", default=None)
    #args = parser.parse_args()
    print(Instagram.scrap(un,"chrome"))

    doc = docx.Document()
  
    # Add a Title to the document
    doc.add_heading('Details', 0)
    data = json.loads(Instagram.scrap(un))
    print(data)
    print(type(data))
    print(len(data))
    table = doc.add_table(rows=1, cols=2)
  
    # Adding heading in the 1st row of the table
    row = table.rows[0].cells
    row[0].text = 'Id'
    row[1].text = 'Value'
    print("hello")
    temp = data
    row = table.add_row().cells
    row[0].text = "Username"
    row[1].text = str(temp['username'])
    print(row[0].text)
    print(row[1].text)
    row = table.add_row().cells
    row[0].text = str("Biography")
    row[1].text = str(temp['biography'])
    row = table.add_row().cells
    row[0].text = str("Follower count")
    row[1].text = str(temp['edge_followed_by']['count'])
    row = table.add_row().cells
    row[0].text = str("Following count")
    row[1].text = str(temp['edge_follow']['count'])
    row = table.add_row().cells
    row[0].text = str("Full name")
    row[1].text = str(temp['full_name'])
    row = table.add_row().cells
    row[0].text = str("Business account")
    row[1].text = str(temp['is_business_account'])
    row = table.add_row().cells
    row[0].text = str("Professional account")
    row[1].text = str(temp['is_professional_account'])
    row = table.add_row().cells
    row[0].text = str("Joined recently")
    row[1].text = str(temp['is_joined_recently'])
    row = table.add_row().cells
    row[0].text = str("Guardian ID")
    row[1].text = str(temp['guardian_id'])
    row = table.add_row().cells
    row[0].text = str("fbid")
    row[1].text = str(temp['fbid'])
    row = table.add_row().cells
    row[0].text = str("Business address")
    row[1].text = str(temp['business_address_json'])
    row = table.add_row().cells
    row[0].text = str("Business contacts")
    row[1].text = str(temp['business_contact_method'])
    row = table.add_row().cells
    row[0].text = str("Business email")
    row[1].text = str(temp['business_email'])
    row = table.add_row().cells
    row[0].text = str("Business phone number")
    row[1].text = str(temp['business_phone_number'])
    row = table.add_row().cells
    row[0].text = str("Business category")
    row[1].text = str(temp['business_category_name'])
    row = table.add_row().cells
    row[0].text = str("Private account")
    row[1].text = str(temp['is_private'])
    row = table.add_row().cells
    row[0].text = str("Verified account")
    row[1].text = str(temp['is_verified'])
    row = table.add_row().cells
    row[0].text = str("Profile pic URL")
    row[1].text = str(temp['profile_pic_url_hd'])
    row = table.add_row().cells
    row[0].text = str("Connected to FB page")
    row[1].text = str(temp['connected_fb_page'])
    print("hello")
    row_count = len(table.rows)
    col_count = len(table.columns)
    print(row_count)
    print(col_count)
    doc.save('instagram.docx')
    print("hello")


# last updated on 1st January, 2023
