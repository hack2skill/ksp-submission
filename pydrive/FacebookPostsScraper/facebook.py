from FacebookPostsScraper.FacebookPS import FacebookPostsScraper as Fbps
from pprint import pprint as pp
from facebook_scraper import get_profile
import docx
def sacebook(un):
    doc=docx.Document()
    # Enter your Facebook email and password
    email = 'policehackathon23@gmail.com'
    password = 'Police@123'
    print("Helloworld1")
    # Instantiate an object
    fps = Fbps(email, password)
    print("Helloworld2")
    print(get_profile(un,cookies='FacebookPostsScraper/facebook.com_cookies.txt'))
    profile= get_profile(un,cookies='FacebookPostsScraper/facebook.com_cookies.txt')
    # Example with single profile
    single_profile = 'https://www.facebook.com/{}'.format(profile['Contact Info']['Facebook'])
    data = fps.get_posts_from_profile(single_profile)
    #s=pp(data)
    p=profile['top_post']
    doc.add_heading('Details',0)
    table=doc.add_table(rows=1, cols=2)
    row = table.rows[0].cells
    row[0].text='ID'
    row[1].text = 'Value'
    row = table.add_row().cells
    row[0].text = 'User ID'
    row[1].text = p['user_id']
    row = table.add_row().cells
    row[0].text = 'Username'
    row[1].text = p['username']
    row = table.add_row().cells
    row[0].text = 'User URL'
    row[1].text = p['user_url']
    row = table.add_row().cells
    row[0].text = 'Cover Photo'
    row[1].text = profile['cover_photo']
    row = table.add_row().cells
    row[0].text = 'Profile Picture'
    row[1].text = profile['profile_picture']
    row = table.add_row().cells
    row[0].text = 'Name'
    row[1].text = profile['Name']
    row = table.add_row().cells
    row[0].text = 'Category'
    row[1].text = profile['Category']
    table1=doc.add_table(rows=1, cols=4)
    row = table1.add_row().cells
    row[0].text = 'Date published'
    row[1].text = 'Description'
    row[2].text = 'Images'
    row[3].text = 'Post URL'
    for d in data:
        row=table1.add_row().cells
        row[0].text = d['published']
        row[1].text = d['description']
        row[2].text = d['images']
        row[3].text = d['post_url']

    doc.save("facebook.docx")

def scrapFacebook(un):
    sacebook(un)